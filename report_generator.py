import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


def format_document(document):
    """Applies professional formatting to the document."""
    # Set styles for headings and paragraphs
    styles = document.styles

    # Title Style
    title_style = styles['Heading 1']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(20)
    title_style.font.bold = True

    # Heading Style
    heading_style = styles['Heading 2']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(18)
    heading_style.font.bold = True

    subheading_style = styles['Heading 3']
    subheading_style.font.name = 'Arial'
    subheading_style.font.size = Pt(16)
    subheading_style.font.bold = True

    # Normal Paragraph Style
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)

def add_horizontal_line(document):
    """Adds a horizontal line using a paragraph of underscores."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run("_" * 60)  # 60 underscores to create a line
    run.bold = True  # Makes the line bold for better visibility
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the line

def generate_reports(research_data, use_cases, datasets, company, industry):
    print("Generating reports...")
    print("Research Data:", research_data)
    print("Use Cases:", use_cases)
    print("Datasets:", datasets)

    # Generate Market Research Report
    if research_data and isinstance(research_data, list):
        try:
            research_df = pd.DataFrame(research_data)
            research_df.to_excel("Market_Research_Report.xlsx", index=False)
            print("Market Research Report created successfully!")
        except Exception as e:
            print("Error creating Market Research Report:", e)
    else:
        print("No valid research data to generate Market Research Report.")

    # Generate Use Case and Feasibility Report
    if use_cases and isinstance(use_cases, list):
        try:
            document = Document()
            format_document(document)  # Apply formatting

            # Title
            title = document.add_heading(level=1)
            run = title.add_run(f"Use Case Feasibility Report\nCompany: {company}\nIndustry: {industry}")
            run.font.bold = True
            run.font.size = Pt(20)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph("\n")  # Add a spacer line

            # Add Use Cases
            document.add_heading("Innovative Use Cases:", level=2)
            for use_case in use_cases:
                document.add_heading(use_case.get("Use Case", "Untitled"), level=3)

                # Add Description
                description = document.add_paragraph()
                description.add_run("Description: ").bold = True
                description.add_run(use_case.get("Description", "No description provided."))

                # Add Benefits
                benefits_title = document.add_paragraph()
                benefits_title.add_run("Benefits:").bold = True

                benefits = use_case.get("Benefits", [])
                if isinstance(benefits, list):
                    for idx, benefit in enumerate(benefits, start=0):
                        if idx != 0:
                            document.add_paragraph(f"• {idx}. {benefit}")
                else:
                    document.add_paragraph(f"• {benefits}")

                # Add a horizontal line for separation
                add_horizontal_line(document)

            # Add Datasets
            document.add_heading("Datasets", level=2)
            if datasets and isinstance(datasets, list):
                for dataset in datasets:
                    document.add_paragraph(f"Platform: {dataset.get('Platform', 'Unknown')}")
                    document.add_paragraph(f"URL: {dataset.get('URL', 'No URL available')}")
                    #document.add_paragraph("\n")  # Add a spacer line

            # Save the report
            document.save("Use_Case_Feasibility_Report.docx")
            print("Use Case Feasibility Report created successfully in DOC format!")
        except Exception as e:
            print("Error creating Use Case Feasibility Report:", e)
    else:
        print("No valid use case or dataset data to generate Use Case Feasibility Report.")